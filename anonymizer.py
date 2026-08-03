#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Обезличиватель документов
Поддерживает: .docx, .xlsx, .pdf, .txt, .html, .csv
Сохраняет обезличенные копии рядом с оригиналом + файл ключей
"""

import re
import sys
import json
import unicodedata
import os
from pathlib import Path
from datetime import datetime
from copy import deepcopy


# ══════════════════════════════════════════════════════════════
#  ПАТТЕРНЫ
# ══════════════════════════════════════════════════════════════

# Граница предложения для адресных паттернов: точка, за которой идёт
# заглавная буква / плейсхолдер / конец текста. Точки адресных сокращений
# («г.», «ул.», «д.», «корп.» …) и инициалов («Б. Полянка») границей
# не считаются — их отсекают негативные lookbehind'ы.
# Нужна, чтобы захват адреса не перепрыгивал в следующее предложение
# («…д. 7. Контакт: [ФИО_1]» — плейсхолдер попадал внутрь ключа АДРЕС).
_SENT_END = (
    r'(?<!\bг)(?<!\bул)(?<!\bпр)(?<!\bпл)(?<!\bш)(?<!\bнаб)(?<!\bпер)(?<!\bбульв)'
    r'(?<!\bд)(?<!\bдом)(?<!\bстр)(?<!\bкорп)(?<!\bк)(?<!\bкв)(?<!\bофис)(?<!\bпом)'
    r'(?<!\bэт)(?<!\bобл)(?<!\bпос)(?<!\bп)(?<!\bс)(?<!\bдер)(?<!\bмкр)(?<!\bим)'
    r'(?<!\b[А-ЯЁA-Z])'
    r'\.(?=\s+[А-ЯЁA-Z\[]|\s*$)'
)

# Кусок значения адреса: без запятых и переводов строки, точка допустима
# только внутри сокращений — на границе предложения захват останавливается
_ADDR_CHUNK = r'(?:(?!' + _SENT_END + r')[^\n,])+'

# Части доменного имени. Зоны перечислены закрытым списком: открытый шаблон
# «слово.слово» превратил бы в сайт любое сокращение с точкой.
_DOMEN_BUKVA = r'[A-Za-zА-Яа-яЁё0-9]'
_DOMEN_METKA = _DOMEN_BUKVA + r'[A-Za-zА-Яа-яЁё0-9-]*'
_DOMEN_ZONA = r'(?i:ru|рф|com|net|org|su|io|pro|biz|info|me|by|kz)'

RULES = [
    # ВАЖНО: порядок имеет значение. Сначала специфичные/длинные паттерны,
    # затем общие/короткие. Replacer заменяет только group(1) если она есть,
    # сохраняя префикс ("ИНН поручителя:" и т.п.) нетронутым.

    # ─── Реквизиты с явными метками (защита подписей) ───────────────────────
    {
        "id": "company",
        "label": "КОМПАНИЯ",
        "patterns": [
            # ООО "СЗ "Санино 1" — вложенные одинаковые кавычки
            r'(?:ООО|ОАО|ЗАО|АО|ПАО|НАО|ГУП|МУП|ФГУП|НКО|АНО|КФХ|ПК|ТСЖ|НП|Банк|Фонд)\s+"[А-ЯЁA-Z]{1,8}\s+"[^"\n]+"',
            r'(?:ООО|ОАО|ЗАО|АО|ПАО|НАО|ГУП|МУП|ФГУП|НКО|АНО|КФХ|ПК|ТСЖ|НП|Банк|Фонд)\s+[«"„][\wА-Яа-яёЁ\s\-\.]+[»""»]',
            r'(?:ООО|ОАО|ЗАО|АО|ПАО|НАО|ГУП|МУП|ФГУП|НКО|АНО|КФХ|ПК|ТСЖ|НП)\s+[А-ЯЁ][А-Яа-яёЁ\s\-]{2,30}(?=[\s,\.;)]|$)',
            r'ИП\s+[А-ЯЁ][а-яё]+\s+[А-ЯЁ]\.[А-ЯЁ]\.',
            r'ИП\s+[А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+',
        ]
    },
    {
        "id": "project",
        "label": "ПРОЕКТ",
        "patterns": [
            r'(?:проект|объект|программа|инициатива)\s+[«"„][\wА-Яа-яёЁ\s\-\.]+[»""»]',
        ]
    },

    # ─── Идентификаторы недвижимости ─────────────────────────────────────────
    {
        "id": "egrn",
        "label": "ЕГРН",
        "patterns": [
            # Полная запись регистрации: 77:18:0170408:1873-77/060/2025-1
            r'\b\d{2}:\d{2}:\d{6,7}:\d{1,5}\s*-\s*\d{2}/\d{3}/\d{4}-\d+\b',
        ]
    },
    {
        "id": "kadastr",
        "label": "КАДАСТР",
        "patterns": [
            r'\b\d{2}:\d{2}:\d{6,7}:\d{1,5}\b',
        ]
    },

    # ─── ФИО (с поддержкой падежей) ──────────────────────────────────────────
    {
        "id": "fio",
        "label": "ФИО",
        "patterns": [
            # Фамилия Имя Отчество — отчество с падежными окончаниями
            r'(?<![А-Яа-яёЁ])[А-ЯЁ][а-яё]{1,20}\s+[А-ЯЁ][а-яё]{1,20}\s+[А-ЯЁ][а-яё]*(?:вич|вна|ич|на)(?:а|у|ем|е|ой|ы|ю)?\b',
            r'(?<![А-Яа-яёЁ])[А-ЯЁ][а-яё]{1,20}\s+[А-ЯЁ]\.[А-ЯЁ]\.',
            r'[А-ЯЁ]\.[А-ЯЁ]\.\s*[А-ЯЁ][а-яё]{1,20}(?![а-яёА-ЯЁ])',
        ]
    },

    # ─── Контекстные реквизиты (с префиксом-меткой) ──────────────────────────
    {
        "id": "passport_issued",
        "label": "ВЫДАН",
        "patterns": [
            # Кем выдан паспорт: "выдан: ГУ МВД РОССИИ ПО Г. МОСКВЕ"
            r'(?:выдан\s*:?\s*)((?:ГУ|УМВД|МВД|ОВД|УФМС|ФМС|ОУФМС|ТП|ОТДЕЛЕНИЕ|ОТДЕЛ)[^,;\n]{5,120}?)(?=,\s*(?:[Дд]ата|код\s+подразделения)|\s*$)',
        ]
    },
    {
        "id": "passport",
        "label": "ПАСПОРТ",
        "patterns": [
            # "серия 4524 № 795984" — захватываем серию и номер вместе
            r'(?:серия\s+)(\d{4}\s*(?:№|No|N|no|Но)?\s*\d{6})',
        ]
    },
    {
        "id": "snils",
        "label": "СНИЛС",
        "patterns": [
            r'(?:СНИЛС\s*[№:]?\s*)(\d{3}[-\s]?\d{3}[-\s]?\d{3}[-\s]?\d{2})',
            r'(?:СНИЛС\s*[№:]?\s*)(\d{11})',
        ]
    },
    {
        "id": "podrazd",
        "label": "КОД_ПОДР",
        "patterns": [
            r'(?:код\s+подразделения\s*[:№]?\s*)(\d{3}[-\s]?\d{3})',
        ]
    },
    {
        "id": "date_birth",
        "label": "ДР",
        "patterns": [
            # Дата рождения по контексту
            r'(\d{2}\.\d{2}\.\d{4})(?=\s*(?:года\s+рождения|г\.\s*р\.))',
        ]
    },
    {
        "id": "date_issue",
        "label": "ДАТА_ВЫДАЧИ",
        "patterns": [
            r'(?:[Дд]ата\s+выдачи\s*:?\s*)(\d{2}\.\d{2}\.\d{4})',
        ]
    },
    {
        "id": "place_birth",
        "label": "МЕСТО_РОЖДЕНИЯ",
        "patterns": [
            # Место рождения: ... до "паспорт" или конца строки
            r'(?:место\s+рождения\s*:\s*)([^,\n]{3,200}?)(?=,\s*(?:паспорт|серия)|\s*$)',
        ]
    },
    {
        "id": "address",
        "label": "АДРЕС",
        "patterns": [
            # Все типичные метки адресов; _SENT_END останавливает захват
            # на границе предложения
            r'(?:зарегистрирован[а-яё]*\s+по\s+адресу\s*:\s*)([^\n]+?)(?=,\s*(?:СНИЛС|паспорт|серия|телефон)|\s*\(далее|' + _SENT_END + r'|\s*$)',
            r'(?:адрес\s+регистрации\s*:\s*)([^\n]+?)(?=,\s*(?:СНИЛС|паспорт|серия|телефон)|\s*\(далее|' + _SENT_END + r'|\s*$)',
            # Юридический/почтовый/фактический адрес
            r'(?:(?:юридический|почтовый|фактический|адрес проживания|адрес)\s*(?:адрес)?\s*:\s*)((?:г(?:ород|\.)?\s+)?[А-ЯЁа-яёA-Za-z][^\n]{10,200}?)(?=,\s*(?:ИНН|ОГРН|СНИЛС|паспорт|телефон|расчётный|р/с|БИК)|' + _SENT_END + r'|\s*$)',
            # Адрес без метки: "г Москва, ул Красноказарменная, д 14А к 6, кв 200"
            #
            # Замер 03.08.2026 (ДЗ PEd06) на синтетическом корпусе показал три
            # формы, на которых прежняя редакция молчала, а все три встречаются
            # в живом входе Аси:
            #   «проживающей по адресу г. Тверь, пр-т Победы, д. 3, кв. 41»
            #        — тип улицы «пр-т» не был перечислен (а меточное правило
            #          выше требует двоеточия, которого в шапке заявления нет);
            #   «г. Казань, ул. Лесная, 14» (визитка)
            #        — номер дома без «д.», а группа требовала сокращение;
            #   «г. Нижний Новгород, ул. Большая Покровская, д. 18»
            #        — город из двух слов, захват обрывался на «Нижний».
            # Улица из двух слов работала и раньше: _ADDR_CHUNK идёт до запятой.
            r'\bг(?:ород|\.)?\s+[А-ЯЁ][а-яё]+(?:-[А-Яа-яЁё]+)*(?:\s+[А-ЯЁ][а-яё]+)?'
            r'\s*,\s*(?:пр-т|просп|бульв|б-р|мкр|наб|пер|ул|пр|ш|пл)\.?\s+[А-ЯЁа-яё]' + _ADDR_CHUNK +
            # Хвост адреса: либо сокращение с номером, либо ГОЛЫЙ номер дома.
            # Голая цифра допускается только здесь, после распознанной улицы:
            # разрешить её раньше значит начать хватать любые перечисления.
            r'(?:\s*,\s*(?:(?:д|дом|стр|корп|к|кв|офис|оф|пом|эт)\.?\s*' + _ADDR_CHUNK + r'|\d+[А-Яа-яЁё]?))+',
        ]
    },

    # ─── Финансовые/налоговые идентификаторы ─────────────────────────────────
    # СЧЁТ — 20 цифр, длиннее всех остальных, идёт первым
    {
        "id": "account",
        "label": "СЧЁТ",
        "patterns": [
            r'(?:р/с|к/с|расч[её]тный\s+счёт|расч[её]тный\s+счет|счёт|счет)\s*[№:\s]*(\d{20})',
            r'\b\d{20}\b',  # любые 20 цифр подряд
        ]
    },
    # ОГРН — 13 цифр (юрлицо) или 15 цифр (ИП)
    {
        "id": "ogrn",
        "label": "ОГРН",
        "patterns": [
            r'(?:ОГРН[ИП]?\s*[:/№]?\s*)(\d{13,15})',
            r'\b\d{15}\b',  # ОГРНИП
            r'\b\d{13}\b',  # ОГРН
        ]
    },
    # ИНН физлица — 12 цифр (идёт ДО ИНН юрлица, потому что 10 ⊂ 12)
    {
        "id": "inn_fl",
        "label": "ИНН_ФЛ",
        "patterns": [
            r'(?:ИНН(?:[\s\w\(\)]*?)?\s*[:/№]?\s*)(\d{12})(?!\d)',
            r'\b\d{12}\b',  # любые 12 цифр = ИНН ФЛ
        ]
    },
    # ИНН юрлица — 10 цифр
    {
        "id": "inn_ul",
        "label": "ИНН_ЮЛ",
        "patterns": [
            r'(?:ИНН(?:[\s\w\(\)]*?)?\s*[:/№]?\s*)(\d{10})(?!\d)',
            r'\b\d{10}\b',  # любые 10 цифр
        ]
    },
    # КПП — 9 цифр (с обязательной меткой, чтобы не путать с БИК)
    {
        "id": "kpp",
        "label": "КПП",
        "patterns": [
            r'(?:КПП\s*[:/№]?\s*)(\d{9})',
        ]
    },
    # БИК — 9 цифр, начинается с 04 для российских банков
    {
        "id": "bik",
        "label": "БИК",
        "patterns": [
            r'(?:БИК(?:[\s\w]*?)?\s*[:/№]?\s*)(\d{9})',
            r'\b04\d{7}\b',  # БИК всегда начинается с 04
        ]
    },
    {
        "id": "swift",
        "label": "SWIFT",
        "patterns": [
            r'(?:SWIFT|свифт)\s*[:/]?\s*([A-Z]{6}[A-Z0-9]{2,5})',
        ]
    },

    # ─── Контакты ───────────────────────────────────────────────────────────
    {
        "id": "phone",
        "label": "ТЕЛЕФОН",
        "patterns": [
            r'(?:\+7|8)[\s\-]?\(?\d{3}\)?[\s\-]?\d{3}[\s\-]?\d{2}[\s\-]?\d{2}',
            r'\b[78]\d{10}\b',  # 11 цифр начинающиеся с 7 или 8 (слитно)
        ]
    },
    {
        "id": "email",
        "label": "EMAIL",
        "patterns": [
            r'[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}',
        ]
    },
    # ─── Сайт: ТОЛЬКО голый домен, ссылку не трогаем ────────────────────────
    #
    # Домен раскрывает компанию не хуже её названия, а для NER это не ПДн, —
    # пробел был записан в HANDOFF с 16.07 и подтверждён замером 03.08 (ДЗ PEd06).
    # Закрыт УЗКИМ вариантом по решению Дмитрия: маскируется только упоминание
    # сайта, ссылка на ресурс остаётся целой.
    #
    # ⚠️ Различитель — НАЛИЧИЕ ПУТИ, а не схемы. Эталонная ссылка в тестах
    # проекта записана без «https://» (`us06web.zoom.us/j/84213?pwd=…`), и
    # критерий «нет схемы» её бы не спас. Инвариант `test_link_survives_the_shield`
    # держится тем, что за путём/запросом/якорем шаблон не идёт.
    # Ссылки — главная ценность события у Аси, поэтому здесь консервативно:
    # что-то похожее на ссылку лучше пропустить, чем испортить.
    #
    # ⚠️ Зона `москва` в список НЕ входит сознательно: «г.Москва» без пробела —
    # обычная запись адреса, и она стала бы доменом. Собственные домены Дмитрия
    # (`финдир.москва`) в маскировке не нуждаются — это не чужие ПДн.
    #
    # Идёт ПОСЛЕ email: к этому месту почта уже стала плейсхолдером, и лезть
    # внутрь неё шаблону не приходится.
    {
        "id": "website",
        "label": "САЙТ",
        "patterns": [
            # слева: не середина домена, не хвост почты, не кусок пути URL
            r'(?<![\w@./-])'
            r'(?:www\.)?' + _DOMEN_METKA + r'(?:\.' + _DOMEN_METKA + r')*'
            r'\.' + _DOMEN_ZONA +
            # справа: ни пути, ни запроса, ни якоря — иначе это ссылка.
            # Точка допускается только как конец предложения (за ней не метка).
            r'(?![\w/?#-])(?!\.' + _DOMEN_BUKVA + r')',
        ]
    },

    # ─── Опциональные (выключены по умолчанию) ──────────────────────────────
    {
        "id": "amount",
        "label": "СУММА",
        "patterns": [
            r'\d[\d\s]*(?:[,\.]\d+)?\s*(?:₽|руб(?:лей|ля|\.)?|\$|€|USD|EUR|RUB)\b',
        ]
    },
    {
        "id": "date",
        "label": "ДАТА",
        "patterns": [
            r'\d{2}\.\d{2}\.\d{4}',
            r'\d{1,2}\s+(?:января|февраля|марта|апреля|мая|июня|июля|августа|сентября|октября|ноября|декабря)\s+\d{4}',
        ]
    },
]

# По умолчанию включены все, кроме сумм и обычных дат (даты рождения/выдачи всегда включены)
ENABLED_BY_DEFAULT = {"company", "project", "egrn", "kadastr", "fio",
                      "passport_issued", "passport", "snils", "podrazd", "date_birth", "date_issue",
                      "place_birth", "address",
                      "account", "ogrn", "inn_ul", "inn_fl", "kpp", "bik", "swift",
                      "phone", "email", "website"}


# ══════════════════════════════════════════════════════════════
#  ЛИЧНЫЙ СЛОВАРЬ (нулевой проход)
# ══════════════════════════════════════════════════════════════

# Файл рядом со скриптом: одна запись на строку, # — комментарий.
# Формат: «Название» или «Название | МЕТКА» (по умолчанию КОМПАНИЯ).
# Закрывает то, что не ловят ни regex, ни NER: названия клиентов
# латиницей без ОПФ (FinClever), внутренние имена проектов.
# Файл в .gitignore — клиентские данные не попадают в репозиторий.
DICT_FILENAME = "словарь_клиентов.txt"

# Метка → type_id существующих правил, чтобы нумерация плейсхолдеров
# не пересекалась между слоями ([КОМПАНИЯ_1] всегда уникален)
DICT_LABEL_TO_ID = {"КОМПАНИЯ": "company", "ФИО": "fio", "ПРОЕКТ": "project"}


def _dict_pattern(entry):
    """Паттерн для записи словаря: границы слова, без учёта регистра.
    Каждому русскому слову записи разрешается падежное окончание
    («Ромашка Трейд» найдёт и «Ромашкой Трейд», и «Ромашке Трейд»)."""
    words = []
    for word in entry.split():
        if re.search(r'[А-Яа-яёЁ]$', word) and len(word) >= 4:
            # Срезаем окончание: двухбуквенное у прилагательных (-ая/-ий/-ое),
            # однобуквенное у существительных (-а/-я/-ь)
            stem = re.sub(r'(?:ая|яя|ий|ый|ое|ее|[аяьйео])$', '', word)
            words.append(re.escape(stem) + r'[а-яё]{0,3}')
        else:
            words.append(re.escape(word))
    body = r'\s+'.join(words)
    return re.compile(
        r'(?<![А-Яа-яёЁA-Za-z0-9_])' + body + r'(?![А-Яа-яёЁA-Za-z0-9_])',
        re.IGNORECASE)


def load_custom_dict(path=None):
    """Возвращает ([(pattern, type_id, label), ...], путь_к_файлу).
    Длинные записи первыми, чтобы «Ромашка Трейд» не перекрылась «Ромашкой»."""
    p = Path(path) if path else Path(__file__).resolve().parent / DICT_FILENAME
    if not p.exists():
        return [], p
    entries = []
    for line in p.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if not line or line.startswith("#"):
            continue
        if "|" in line:
            value, label = [s.strip() for s in line.split("|", 1)]
            label = label.upper() or "КОМПАНИЯ"
        else:
            value, label = line, "КОМПАНИЯ"
        if len(value) >= 3:
            entries.append((value, label))
    entries.sort(key=lambda e: -len(e[0]))
    return [(_dict_pattern(v), DICT_LABEL_TO_ID.get(l, l.lower()), l)
            for v, l in entries], p


# ══════════════════════════════════════════════════════════════
#  NER-СЛОЙ (Natasha)
# ══════════════════════════════════════════════════════════════

# Типы сущностей Natasha → (type_id, label) анонимайзера.
# PER и ORG используют те же id, что и regex-правила, поэтому
# нумерация плейсхолдеров и дедупликация сквозные для обоих слоёв.
NER_LABEL_MAP = {
    "PER": ("fio", "ФИО"),
    "ORG": ("company", "КОМПАНИЯ"),
    "LOC": ("ner_loc", "ЛОКАЦИЯ"),
}

# LOC по умолчанию выключен: упоминания городов/регионов обычно нужны
# для анализа, а адреса ловит regex-правило address
NER_TYPES_BY_DEFAULT = ("PER", "ORG")

# Канцелярские роли, которые NER-модель путает с ФИО в шапках документов
# («Заёмщик: ...», «Поручитель: ...»)
NER_ROLE_STOPWORDS = {
    "заёмщик", "заемщик", "поручитель", "залогодатель", "кредитор",
    "арендатор", "арендодатель", "покупатель", "продавец",
    "исполнитель", "заказчик", "подрядчик", "клиент", "агент",
    "директор", "руководитель", "бухгалтер", "стороны", "сторона",
    "договор", "приложение", "банк",

    # ─── Деловые события и документы (добавлено 03.08.2026) ───────────────
    # Находка Г из `plans/2026-07-16-nahodki-boevogo-progona.md` («План-фактный»
    # → [ФИО_2]) была отложена с триггером «если карточки начнут врать».
    # Триггер сработал: живая проверка после выката 03.08 дала
    # «Созвон завтра» → «[ФИО_1] завтра». Это типичный вход Аси, и после
    # маскировки LLM извлекает событие, у которого название стало плейсхолдером.
    #
    # ⚠️ Почему списком, а не правилом. Три принципиальных способа замерены
    # и отвергнуты (разбор — в HANDOFF, раздел «Что сделано 03.08.2026»):
    # разборщик имён natasha требует словарей pymorphy2, которых нет;
    # нейроморфология НЕ различает («Созвон» и «Литвинов» одинаково
    # Animacy=Anim, POS=PROPN); словарь pymorphy3 различает, но правило
    # «неодушевлённое нарицательное» перестало бы маскировать 14 настоящих
    # фамилий из 15 проверенных (Кисель, Борщ, Топор, Ключ, Зима, Сокол…),
    # а утечка дороже переусердствования.
    #
    # ⚠️ Каждое слово здесь обязано быть проверено словарём на то, что оно НЕ
    # фамилия и не имя. Это не пожелание: проверку делает машинa —
    # `test_stopwords_are_not_surnames`. При добавлении слов гонять её.
    # Кандидат «вебинар» этой проверкой уже отклонён.
    "созвон", "планёрка", "планерка", "приёмка", "приемка", "отгрузка",
    "оплата", "встреча", "совещание", "отчёт", "отчет", "смета", "заявка",
    "поставка", "доставка", "сверка", "ревизия", "инвентаризация",
    "конференция", "эфир", "интервью", "собеседование", "переговоры",
    "презентация", "выгрузка", "отправка", "напоминание", "задача",
    "дедлайн", "акт", "накладная", "спецификация", "тендер", "аукцион",
    "аванс", "предоплата", "отсрочка", "отпуск", "командировка",
    "обучение", "тренинг", "звонок", "письмо", "напоминалка", "запись",
    "подпись", "согласование", "вебинар",

    # ─── Финансовые термины и метрики (добавлено 03.08.2026) ──────────────
    # Прогон по 7 РЕАЛЬНЫМ документам Дмитрия (кредитная политика, кейсовые
    # задачи, методики) дал 68 масок, из них 16 — вот эти аббревиатуры,
    # уехавшие в [КОМПАНИЯ]. Документ, где «МСФО» и «WACC» заменены
    # плейсхолдерами, для LLM бесполезен, а именно такие документы —
    # рабочий материал finsvc.
    "мсфо", "рсбу", "цфо", "осв", "ддс", "опиу", "ндс", "ндфл", "усн",
    "осно", "енвд", "пбу", "гаап", "квэд", "оквэд",
    "cfo", "cfa", "ceo", "coo", "cto", "cio", "wacc", "dscr", "ebitda",
    "ebit", "capex", "opex", "roi", "roe", "roa", "irr", "npv", "ccc",
    "dio", "dso", "dpo", "abc", "xyz", "kpi", "erp", "crm", "llm",
    "acca", "cima", "ifrs", "gaap", "ltv", "cac", "arr", "mrr", "bi",
    "sql", "api", "pdf", "ocr", "rfm",

    # ─── Заголовочные слова из тех же живых документов ────────────────────
    # Класс «обычное слово с большой буквы → [ФИО]»: 24 случая из 68.
    # Глагольные формы («Рассчитайте», «Ответьте») списком не покрыть —
    # они отсекаются отдельным признаком, см. `_ner_slovo_ne_imya`.
    "глоссарий", "годовой", "дашборд", "дашборда", "дилерские", "пассив",
    "актив", "ответ", "скидка", "объём", "объем", "премия", "выручка",
    "прибыль", "убыток", "баланс", "ежеквартально", "ежемесячно",
    "обязательно", "инвест", "кейсовая", "итого", "прочее", "примечание",
}

# Маркер списка в начале строки — типовая вёрстка Outlook/Word
# («·        Литвинов Павел, …»). Символы «·» и «•» слепят Natasha:
# NER не видит ФИО в такой строке (утечка на приёмке 07.07.2026).
# Перед NER-проходом маркер заменяется пробелом ТОЙ ЖЕ длины, поэтому
# спаны NER остаются валидными в исходном тексте — обратный маппинг
# смещений не нужен. Дефис, тире и звёздочка нейтрализуются только
# перед пробельным символом, чтобы не задеть осмысленные начала строк.
_NER_LINE_MARKER = re.compile(r'(?m)^(\s*)(?:[·•‣◦▪▫●○■□∙]|[*—–-](?=\s))')

# Латинские буквы, НЕОТЛИЧИМЫЕ от кириллических по начертанию. Тот же класс
# утечки, что маркер списка выше: слово выглядит русским, а NER его не видит.
# Живой случай 16.07.2026: OCR прочитал логотип на визитке латиницей — буквы
# A U+0041, T U+0054, O U+004F, H U+0048 вместо кириллических. На экране не
# отличить, и человек в петле такое не поймает — чинить можно только здесь.
# В карте только настоящие двойники: строчные «м/m», «т/t», «в/b» — не двойники.
_LATIN_TO_CYRILLIC = {
    'A': 'А', 'B': 'В', 'E': 'Е', 'K': 'К', 'M': 'М', 'H': 'Н',
    'O': 'О', 'P': 'Р', 'C': 'С', 'T': 'Т', 'X': 'Х',
    'a': 'а', 'e': 'е', 'o': 'о', 'p': 'р', 'c': 'с', 'x': 'х', 'y': 'у',
}
# Ссылки, почты, коды и всё с цифрами — мимо: там латиница настоящая, а лишний
# спан NER увёл бы ссылку в плейсхолдер (замер 16.07: в токене пароля Zoom
# одиночная «a» формально «неоднозначна»).
_NER_UNTOUCHABLE = re.compile(r'\S*\w[./@:=?&\\]\w\S*|\S*\d\S*')
_NER_WORD = re.compile(r'[^\W\d_]+')


def _restore_alphabet(match):
    word = match.group(0)
    if len(word) < 2:
        return word  # одиночная буква «неоднозначна» всегда
    latin = [ch for ch in word if ch.isascii()]
    if not latin:
        return word  # уже кириллица
    if any(ch not in _LATIN_TO_CYRILLIC for ch in latin):
        # Есть латинская буква без двойника — слово латинское по-настоящему
        # («Power»: у «w» двойника нет, значит это не искажённая кириллица)
        return word
    return ''.join(_LATIN_TO_CYRILLIC.get(ch, ch) for ch in word)


def _ner_view(text):
    """Копия текста для NER-прохода: маркеры списка заменены пробелами,
    латиница-двойники возвращена в кириллицу.

    Длина текста не меняется — индексы совпадают с исходным, поэтому спаны
    NER валидны в оригинале, а сам оригинал остаётся нетронутым: наружу
    уходит исходное написание, подменяется только то, что NER читает.
    """
    view = _NER_LINE_MARKER.sub(lambda m: m.group(1) + ' ', text)
    parts = []
    last = 0
    for chunk in _NER_UNTOUCHABLE.finditer(view):
        parts.append(_NER_WORD.sub(_restore_alphabet, view[last:chunk.start()]))
        parts.append(chunk.group(0))
        last = chunk.end()
    parts.append(_NER_WORD.sub(_restore_alphabet, view[last:]))
    return ''.join(parts)


class SurnameDictLayer:
    """Третий проход: словарь фамилий (pymorphy3) добирает то, что пропустил NER.

    Зачем. 03.08.2026 нашлось, что NER пропускает фамилии **спорадически**:
    «Морозова прислала смету» проходит насквозь, «Морозов прислал смету» —
    маскируется. Замер на 36 сочетаниях (6 фамилий × пол × падеж × позиция) дал
    3 пропуска — это не правило вроде «женская в начале фразы», а промахи модели.
    Пропуск = утечка, поэтому нужен добор.

    ⚠️ Правило намеренно узкое: слово с большой буквы, которое словарь знает как
    ФАМИЛИЮ (`Surn`) и НЕ знает как имя или отчество (`Name`/`Patr`).
    Сужение до «не имя» — не украшение, а условие совместимости с решением
    от 16.07.2026 «словарь имён отменён замером»: широкое правило по `Surn`
    маскирует «Роман с продолжением», потому что фамилия Роман существует.
    Проверено: «Вера», «Роман», «Любовь», «Надежда» остаются нетронутыми.

    ⚠️ Цена названа: слово, которое одновременно фамилия и обычное слово,
    маскируется лишнего — из 21 проверенного делового слова таким оказалось
    одно, «Мороз». Это переусердствование, а не утечка, то есть безопасная
    сторона размена.

    Слой необязателен: без pymorphy3 работает как раньше (то есть пропуск
    остаётся) — зависимость мягкая, как и natasha.
    """

    # Слово с большой буквы целиком: не часть плейсхолдера, не середина слова.
    SLOVO = re.compile(r'(?<![\w\[])[А-ЯЁ][а-яё]+(?![\w\]])')

    def __init__(self):
        import pymorphy3
        self._morph = pymorphy3.MorphAnalyzer()

    LICHNOE = {"Name", "Surn", "Patr"}
    GLAGOL = {"VERB", "INFN", "PRTF", "PRTS", "GRND"}

    def is_personal_name(self, word):
        """Слово — личное имя по словарю. Объединение двух признаков.

        ⚠️ Ни один по отдельности не годится, замерено 03.08.2026:

        A. «фамилия и НЕ имя» — берёт «Литвинов», но теряет имена и отчества
           («Марина», «Ксения», «Сергеевич»), а именно имя стояло первой строкой
           на реальной визитке и уходило наружу.
        B. «все разборы личные» — берёт имена и отчества, но теряет «Литвинов»
           (у него есть и другое чтение).

        Объединение берёт 12 форм из 14 проверенных и не задевает ни одного
        делового слова из 17. Пропускает «Максим» и «Петров» — у обоих есть
        нарицательное чтение; их берёт NER в контексте, слой лишь страховка.
        """
        g = [set(p.tag.grammemes) for p in self._morph.parse(word)]
        if not g:
            return False
        familiya_ne_imya = (any("Surn" in r for r in g)
                            and not any(("Name" in r) or ("Patr" in r) for r in g))
        vse_lichnye = all(self.LICHNOE & r for r in g)
        return familiya_ne_imya or vse_lichnye

    def is_verb_form(self, word):
        """Все разборы — глагольные и ни одного личного.

        Нужно для обратной задачи: NER принимает повелительные формы из
        методичек («Рассчитайте», «Ответьте», «Составьте») за ФИО. Списком
        такое не покрыть — форм слишком много. Замер: 7 глаголов из 7 отсеяны,
        10 имён и фамилий из 10 не задеты.
        """
        g = [set(p.tag.grammemes) for p in self._morph.parse(word)]
        return bool(g) and all((self.GLAGOL & r) and not (self.LICHNOE & r) for r in g)


class NerLayer:
    """Второй проход: нейросетевой NER (Natasha/Slovnet) добирает ФИО и
    организации, которые не покрыты regex-шаблонами — одиночные фамилии,
    компании без организационно-правовой формы, нестандартные написания.
    Regex-слой всегда идёт первым, чтобы структурные реквизиты
    (ИНН, СНИЛС, счета) не достались NER."""

    def __init__(self, types=NER_TYPES_BY_DEFAULT):
        from natasha import Segmenter, NewsEmbedding, NewsNERTagger, Doc
        self._segmenter = Segmenter()
        self._tagger = NewsNERTagger(NewsEmbedding())
        self._Doc = Doc
        self.types = set(types)

    def spans(self, text):
        """Возвращает [(start, stop, type_id, label, value), ...]"""
        doc = self._Doc(text)
        doc.segment(self._segmenter)
        doc.tag_ner(self._tagger)
        out = []
        for span in doc.spans:
            if span.type not in self.types:
                continue
            type_id, label = NER_LABEL_MAP[span.type]
            out.append((span.start, span.stop, type_id, label, text[span.start:span.stop]))
        return out


class Anonymizer:
    def __init__(self, enabled_types=None, use_ner=True, use_dict=True, dict_path=None):
        self.enabled = enabled_types or ENABLED_BY_DEFAULT
        self.keys = {}        # placeholder → original
        self.counters = {}    # type_id → count
        self._seen = {}       # original → placeholder (для дедупликации)
        self.ner = None
        self.ner_status = "выключен"
        if use_ner:
            try:
                self.ner = NerLayer()
                self.ner_status = "включён (Natasha)"
            except ImportError:
                self.ner_status = "недоступен — pip3 install natasha"
        # Словарный добор фамилий — тоже мягкая зависимость: нет pymorphy3,
        # значит слой молчит и поведение прежнее (с известным пропуском).
        self.surnames = None
        self.surnames_status = "выключен"
        if use_ner:
            try:
                self.surnames = SurnameDictLayer()
                self.surnames_status = "включён (pymorphy3)"
            except ImportError:
                self.surnames_status = "недоступен — pip3 install pymorphy3"
        self.custom_rules = []
        self.dict_status = "выключен"
        if use_dict:
            self.custom_rules, dict_file = load_custom_dict(dict_path)
            if self.custom_rules:
                self.dict_status = f"{len(self.custom_rules)} записей ({dict_file.name})"
            else:
                self.dict_status = f"пусто — заполните {dict_file.name}"

    def _next_placeholder(self, type_id, label):
        self.counters[type_id] = self.counters.get(type_id, 0) + 1
        return f"[{label}_{self.counters[type_id]}]"

    @staticmethod
    def _preprocess(text):
        """Нормализует текст перед обезличиванием:
        - Склеивает цифры разделённые одиночными пробелами (формат ПД-4).
          Пример: "9 7 2 1 1 8 1 7 8 3" → "9721181783"
        - Для длинных склеек (30/29 цифр) разделяет на ИНН+СЧЁТ или БИК+СЧЁТ.
        """
        # Нормализация Unicode: "и" + combining breve → "й"
        # Критически важно для docx из некоторых редакторов
        text = unicodedata.normalize('NFC', text)

        def collapse(m):
            digits = m.group(0).replace(' ', '')
            n = len(digits)
            # Известные комбинации полей формы ПД-4
            if n == 30:   # ИНН (10) + расчётный счёт (20)
                return digits[:10] + ' ' + digits[10:]
            if n == 29:   # БИК (9) + счёт (20) или корсчёт (20) + лицевой (9)
                return digits[:9] + ' ' + digits[9:]
            if n == 32:   # ИНН (12) + счёт (20)
                return digits[:12] + ' ' + digits[12:]
            return digits
        # Минимум 6 цифр разделённых одиночными пробелами
        return re.sub(r'(?:\d\s){5,}\d', collapse, text)

    def replace(self, text):
        """Заменяет конфиденциальные данные в тексте на плейсхолдеры."""
        if not text:
            return text

        # Нормализация: склеиваем цифры с пробелами (ПД-4 формат)
        result = self._preprocess(text)

        # Нулевой проход: личный словарь — приоритет над regex и NER
        for pat, type_id, label in self.custom_rules:

            def dict_replacer(m, type_id=type_id, label=label):
                original = m.group(0)
                key = f"{type_id}::{original}"
                if key in self._seen:
                    return self._seen[key]
                ph = self._next_placeholder(type_id, label)
                self.keys[ph] = {"original": original, "type": label}
                self._seen[key] = ph
                return ph

            result = pat.sub(dict_replacer, result)

        for rule in RULES:
            if rule["id"] not in self.enabled:
                continue
            for pat_str in rule["patterns"]:
                pat = re.compile(pat_str, re.IGNORECASE if rule["id"] in {"email","amount","date","project","swift"} else 0)

                def replacer(m, rule=rule):
                    # Если паттерн имеет группу захвата — заменяем только её,
                    # префикс ("ИНН поручителя:", "БИК:" и т.д.) оставляем как есть
                    try:
                        value = m.group(1).strip() if m.lastindex and m.group(1) else None
                    except IndexError:
                        value = None

                    if value:
                        original = value
                        prefix = m.group(0)[:m.start(1) - m.start(0)]
                        suffix = m.group(0)[m.end(1) - m.start(0):]
                    else:
                        original = m.group(0).strip()
                        prefix = ""
                        suffix = ""

                    if not original:
                        return m.group(0)

                    key = f"{rule['id']}::{original}"
                    if key in self._seen:
                        ph = self._seen[key]
                    else:
                        ph = self._next_placeholder(rule["id"], rule["label"])
                        self.keys[ph] = {"original": original, "type": rule["label"]}
                        self._seen[key] = ph

                    return prefix + ph + suffix

                result = pat.sub(replacer, result)

        if self.ner:
            result = self._apply_ner(result)

        return result

    def _apply_ner(self, text):
        """Второй проход по тексту: NER-сущности, пропущенные regex-слоем.
        Замены выполняются справа налево, чтобы не сбить смещения спанов."""
        if len(text) < 4 or not re.search(r'[А-Яа-яёЁ]', text):
            return text

        # NER смотрит на текст с нейтрализованными маркерами списка;
        # длина совпадает с исходным, поэтому спаны применимы к text.
        # Значение берём из исходной строки — какой она уйдёт в вывод.
        for start, stop, type_id, label, _ in reversed(self.ner.spans(_ner_view(text))):
            value = text[start:stop].strip()
            # Не трогаем уже вставленные плейсхолдеры [ФИО_1] и их окрестности
            if re.search(r'[\[\]]', text[max(0, start - 1):stop + 1]):
                continue
            # Отсекаем мусорные спаны: короче 3 символов, без букв,
            # или похожие на внутренность плейсхолдера (КАПС_1)
            if len(value) < 3 or not re.search(r'[А-Яа-яёЁA-Za-z]', value):
                continue
            if re.fullmatch(r'[А-ЯЁA-Z0-9_]+', value) and '_' in value:
                continue
            # Коды документов («КМ-2024-0055»): цифры есть, строчных букв нет
            if re.search(r'\d', value) and not re.search(r'[а-яёa-z]', value):
                continue
            # Одиночное слово-роль из шапки документа — не ФИО
            if value.lower() in NER_ROLE_STOPWORDS:
                continue
            # Повелительная форма из методички — не ФИО («Рассчитайте», «Ответьте»).
            # Отдельным признаком, а не списком: форм слишком много.
            if (self.surnames is not None and len(value.split()) == 1
                    and self.surnames.is_verb_form(value)):
                continue

            key = f"{type_id}::{value}"
            if key in self._seen:
                ph = self._seen[key]
            else:
                ph = self._next_placeholder(type_id, label)
                self.keys[ph] = {"original": value, "type": label}
                self._seen[key] = ph

            text = text[:start] + text[start:stop].replace(value, ph, 1) + text[stop:]

        return self._surname_dict_pass(text)

    def _surname_dict_pass(self, text):
        """Третий проход: добор фамилий, которые NER пропустил (см. SurnameDictLayer).

        Идёт ПОСЛЕДНИМ и только по словам, оставшимся открытыми: regex и NER уже
        превратили всё найденное в плейсхолдеры, а внутрь плейсхолдера шаблон
        не заходит. Слой только ДОБАВЛЯЕТ маскировку и никогда её не снимает —
        направление, безопасное для fail-closed.
        """
        if self.surnames is None or "fio" not in self.enabled:
            return text

        def zamena(m):
            value = m.group(0)
            if value.lower() in NER_ROLE_STOPWORDS:
                return value
            if not self.surnames.is_personal_name(value):
                return value
            key = f"fio::{value}"
            if key in self._seen:
                return self._seen[key]
            ph = self._next_placeholder("fio", "ФИО")
            self.keys[ph] = {"original": value, "type": "ФИО"}
            self._seen[key] = ph
            return ph

        return SurnameDictLayer.SLOVO.sub(zamena, text)

    def restore(self, text):
        """Заменяет плейсхолдеры обратно на оригинальные данные."""
        return restore_text(text, self.keys)

    def stats(self):
        total = sum(self.counters.values())
        return total, dict(self.counters)


# ══════════════════════════════════════════════════════════════
#  ОБРАБОТЧИКИ ФОРМАТОВ
# ══════════════════════════════════════════════════════════════

def process_txt(path_in, path_out, anon):
    text = path_in.read_text(encoding="utf-8", errors="replace")
    result = anon.replace(text)
    path_out.write_text(result, encoding="utf-8")


def process_html(path_in, path_out, anon):
    from html.parser import HTMLParser

    class HtmlAnonymizer(HTMLParser):
        def __init__(self):
            super().__init__(convert_charrefs=False)
            self.output = []

        def handle_starttag(self, tag, attrs):
            attr_str = ""
            for name, val in attrs:
                if val is None:
                    attr_str += f" {name}"
                else:
                    attr_str += f' {name}="{val}"'
            self.output.append(f"<{tag}{attr_str}>")

        def handle_endtag(self, tag):
            self.output.append(f"</{tag}>")

        def handle_data(self, data):
            self.output.append(anon.replace(data))

        def handle_comment(self, data):
            self.output.append(f"<!--{data}-->")

        def handle_decl(self, decl):
            self.output.append(f"<!{decl}>")

        def unknown_decl(self, data):
            self.output.append(f"<![{data}]>")

        def handle_entityref(self, name):
            self.output.append(f"&{name};")

        def handle_charref(self, name):
            self.output.append(f"&#{name};")

    src = path_in.read_text(encoding="utf-8", errors="replace")
    parser = HtmlAnonymizer()
    parser.feed(src)
    path_out.write_text("".join(parser.output), encoding="utf-8")


def process_csv(path_in, path_out, anon):
    import csv
    rows = []
    with open(path_in, encoding="utf-8", errors="replace", newline="") as f:
        reader = csv.reader(f)
        for row in reader:
            rows.append([anon.replace(cell) for cell in row])
    with open(path_out, "w", encoding="utf-8", newline="") as f:
        csv.writer(f).writerows(rows)


def process_docx(path_in, path_out, anon):
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(str(path_in))

    def process_paragraph(para):
        # Собираем полный текст параграфа
        full_text = "".join(run.text for run in para.runs)
        if not full_text.strip():
            return
        replaced = anon.replace(full_text)
        if replaced == full_text:
            return
        # Записываем результат в первый ран, остальные чистим
        if para.runs:
            para.runs[0].text = replaced
            for run in para.runs[1:]:
                run.text = ""

    def process_table(table):
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_paragraph(para)
                for tbl in cell.tables:
                    process_table(tbl)

    # Параграфы основного текста
    for para in doc.paragraphs:
        process_paragraph(para)

    # Таблицы
    for table in doc.tables:
        process_table(table)

    # Колонтитулы
    for section in doc.sections:
        for hdr in [section.header, section.footer]:
            if hdr:
                for para in hdr.paragraphs:
                    process_paragraph(para)

    doc.save(str(path_out))


def process_xlsx(path_in, path_out, anon):
    import openpyxl
    wb = openpyxl.load_workbook(str(path_in))
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    cell.value = anon.replace(cell.value)
    wb.save(str(path_out))


def process_pdf(path_in, path_out, anon):
    """
    PDF: извлекаем текст, обезличиваем, сохраняем как .txt
    (редактирование PDF с сохранением layout не поддерживается без платных инструментов)
    """
    try:
        import pdfplumber
        pages = []
        with pdfplumber.open(str(path_in)) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                pages.append(anon.replace(text))
        # Выходной файл — .md (PDF нельзя редактировать напрямую)
        out_md = path_out.with_suffix(".md")
        sections = [f"## Страница {i+1}\n\n{text}" for i, text in enumerate(pages)]
        out_md.write_text("\n\n---\n\n".join(sections), encoding="utf-8")
        return out_md
    except ImportError:
        raise RuntimeError("Установите pdfplumber: pip install pdfplumber")



# ══════════════════════════════════════════════════════════════
#  РАСШИФРОВКА
# ══════════════════════════════════════════════════════════════

def load_keys(keys_path: Path) -> dict:
    with open(keys_path, "r", encoding="utf-8") as f:
        return json.load(f)


# Лимит проходов расшифровки: защита от original, содержащего
# собственный плейсхолдер (иначе замены не сойдутся никогда)
RESTORE_MAX_PASSES = 10


def restore_text(text: str, keys: dict) -> str:
    """Заменяет плейсхолдеры обратно на оригинальные значения.

    Многопроходно: original может содержать вложенный плейсхолдер —
    например, [АДРЕС_1] разворачивается в текст, внутри которого стоит
    [ФИО_1]. Замены повторяются, пока текст не перестанет меняться.
    """
    for _ in range(RESTORE_MAX_PASSES):
        prev = text
        for ph, val in keys.items():
            text = text.replace(ph, val["original"])
        if text == prev:
            break
    return text


def restore_txt(path_in, path_out, keys):
    text = path_in.read_text(encoding="utf-8", errors="replace")
    path_out.write_text(restore_text(text, keys), encoding="utf-8")


def restore_html(path_in, path_out, keys):
    text = path_in.read_text(encoding="utf-8", errors="replace")
    path_out.write_text(restore_text(text, keys), encoding="utf-8")


def restore_csv(path_in, path_out, keys):
    import csv
    rows = []
    with open(path_in, encoding="utf-8", errors="replace", newline="") as f:
        reader = csv.reader(f)
        for row in reader:
            rows.append([restore_text(cell, keys) for cell in row])
    with open(path_out, "w", encoding="utf-8", newline="") as f:
        csv.writer(f).writerows(rows)


def restore_docx(path_in, path_out, keys):
    from docx import Document

    doc = Document(str(path_in))

    def restore_paragraph(para):
        full_text = "".join(run.text for run in para.runs)
        if not full_text.strip():
            return
        restored = restore_text(full_text, keys)
        if restored == full_text:
            return
        if para.runs:
            para.runs[0].text = restored
            for run in para.runs[1:]:
                run.text = ""

    def restore_table(table):
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    restore_paragraph(para)
                for tbl in cell.tables:
                    restore_table(tbl)

    for para in doc.paragraphs:
        restore_paragraph(para)
    for table in doc.tables:
        restore_table(table)
    for section in doc.sections:
        for hdr in [section.header, section.footer]:
            if hdr:
                for para in hdr.paragraphs:
                    restore_paragraph(para)

    doc.save(str(path_out))


def restore_xlsx(path_in, path_out, keys):
    import openpyxl
    wb = openpyxl.load_workbook(str(path_in))
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    cell.value = restore_text(cell.value, keys)
    wb.save(str(path_out))


RESTORERS = {
    ".txt":  restore_txt,
    ".md":   restore_txt,
    ".html": restore_html,
    ".htm":  restore_html,
    ".csv":  restore_csv,
    ".docx": restore_docx,
    ".xlsx": restore_xlsx,
    ".xls":  restore_xlsx,
}


def restore_file(path_in: Path, keys: dict) -> Path:
    ext = path_in.suffix.lower()
    restorer = RESTORERS.get(ext)
    if restorer is None:
        raise ValueError(f"Формат {ext} не поддерживается для расшифровки")
    # Убираем суффикс _anon если есть, добавляем _restored
    stem = path_in.stem
    if stem.endswith("_anon"):
        stem = stem[:-5]
    path_out = path_in.parent / f"{stem}_restored{path_in.suffix}"
    restorer(path_in, path_out, keys)
    return path_out


def find_keys_file(folder: Path) -> Path | None:
    """Ищет самый свежий файл ключей в папке."""
    candidates = sorted(folder.glob("ключи_*.json"), reverse=True)
    return candidates[0] if candidates else None


def main_restore(files: list, keys_path_str: str = None):
    print()
    print("=" * 60)
    print("  РАСШИФРОВКА ДОКУМЕНТОВ")
    print("=" * 60)
    print()

    # Определяем файл ключей
    if keys_path_str:
        keys_path = Path(keys_path_str)
    else:
        # Ищем автоматически рядом с первым файлом
        folder = Path(files[0]).parent
        keys_path = find_keys_file(folder)
        if keys_path:
            print(f"  Найден файл ключей: {keys_path.name}")
        else:
            print("  ✗  Файл ключей не найден.")
            print("     Укажите явно: python anonymizer.py --restore файл.docx ключи.json")
            print()
            input("Нажмите Enter для выхода...")
            sys.exit(1)

    if not keys_path.exists():
        print(f"  ✗  Файл ключей не найден: {keys_path}")
        input("Нажмите Enter для выхода...")
        sys.exit(1)

    keys = load_keys(keys_path)
    print(f"  Загружено ключей: {len(keys)}")
    print()

    errors = []
    for f_str in files:
        path_in = Path(f_str)
        if not path_in.exists():
            errors.append(f"  ✗  {path_in.name} — файл не найден")
            continue
        try:
            path_out = restore_file(path_in, keys)
            print(f"  ✓  {path_in.name}")
            print(f"       → {path_out.name}")
        except Exception as e:
            errors.append(f"  ✗  {path_in.name} — {e}")

    print()
    if errors:
        print("Ошибки:")
        for e in errors:
            print(e)
        print()

    print("=" * 60)
    input("Готово. Нажмите Enter для закрытия...")

# ══════════════════════════════════════════════════════════════
#  ГЛАВНАЯ ФУНКЦИЯ
# ══════════════════════════════════════════════════════════════

PROCESSORS = {
    ".txt":  process_txt,
    ".md":   process_txt,
    ".html": process_html,
    ".htm":  process_html,
    ".csv":  process_csv,
    ".docx": process_docx,
    ".xlsx": process_xlsx,
    ".xls":  process_xlsx,
    ".pdf":  process_pdf,
}


def anonymize_file(path_in: Path, anon: Anonymizer) -> Path:
    ext = path_in.suffix.lower()
    processor = PROCESSORS.get(ext)
    if processor is None:
        raise ValueError(f"Формат {ext} не поддерживается")

    # Имя выходного файла: original_anon.ext
    path_out = path_in.parent / f"{path_in.stem}_anon{path_in.suffix}"
    actual_out = processor(path_in, path_out, anon)
    return actual_out or path_out


def save_keys(keys: dict, folder: Path) -> Path:
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    keys_path = folder / f"ключи_{ts}.json"
    with open(keys_path, "w", encoding="utf-8") as f:
        json.dump(keys, f, ensure_ascii=False, indent=2)
    return keys_path


def main():
    args = sys.argv[1:]

    # Режим расшифровки: python anonymizer.py --restore файл1 [файл2...] [ключи.json]
    if args and args[0] == "--restore":
        rest = args[1:]
        if not rest:
            print("Укажите файлы для расшифровки.")
            input("Нажмите Enter для выхода...")
            sys.exit(1)
        # Последний аргумент — файл ключей если .json
        if rest[-1].endswith(".json"):
            main_restore(rest[:-1], rest[-1])
        else:
            main_restore(rest)
        return

    use_ner = "--no-ner" not in args
    use_dict = "--no-dict" not in args
    files = [a for a in args if a not in ("--no-ner", "--no-dict")]

    if not files:
        print("=" * 60)
        print("  ОБЕЗЛИЧИВАТЕЛЬ ДОКУМЕНТОВ")
        print("=" * 60)
        print()
        print("Использование:")
        print("  python anonymizer.py файл1.docx файл2.xlsx ...")
        print("  python anonymizer.py --no-ner файл.docx   (без NER-слоя)")
        print("  python anonymizer.py --no-dict файл.docx  (без личного словаря)")
        print()
        print("Поддерживаемые форматы:")
        print("  .docx  .xlsx  .pdf  .txt  .html  .csv")
        print()
        print("Результат:")
        print("  • filename_anon.ext — обезличенная копия")
        print("  • ключи_YYYYMMDD_HHMMSS.json — файл ключей")
        print()
        input("Нажмите Enter для выхода...")
        sys.exit(0)

    anon = Anonymizer(use_ner=use_ner, use_dict=use_dict)
    results = []
    errors = []

    print()
    print("=" * 60)
    print("  ОБЕЗЛИЧИВАТЕЛЬ ДОКУМЕНТОВ")
    print("=" * 60)
    print()
    print(f"  NER-слой: {anon.ner_status}")
    print(f"  Личный словарь: {anon.dict_status}")
    print()

    for f_str in files:
        path_in = Path(f_str)
        if not path_in.exists():
            errors.append(f"  ✗  {path_in.name} — файл не найден")
            continue
        try:
            path_out = anonymize_file(path_in, anon)
            results.append((path_in, path_out))
            print(f"  ✓  {path_in.name}")
            print(f"       → {path_out.name}")
        except Exception as e:
            errors.append(f"  ✗  {path_in.name} — {e}")

    print()

    if errors:
        print("Ошибки:")
        for e in errors:
            print(e)
        print()

    total, counters = anon.stats()

    if total > 0:
        print(f"Найдено и заменено: {total} значений")
        for type_id, count in counters.items():
            rule = next((r for r in RULES if r["id"] == type_id), None)
            label = rule["label"] if rule else type_id.upper()
            print(f"  {label}: {count}")
        print()

        # Сохраняем файл ключей рядом с первым файлом
        folder = Path(files[0]).parent
        keys_path = save_keys(anon.keys, folder)
        print(f"Файл ключей: {keys_path.name}")
        print(f"Папка: {folder}")
    else:
        print("Конфиденциальных данных не найдено.")

    print()
    print("=" * 60)
    input("Готово. Нажмите Enter для закрытия...")


if __name__ == "__main__":
    main()
